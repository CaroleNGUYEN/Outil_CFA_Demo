import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
import warnings
import streamlit as st
import time
import plotly.express as px
import plotly.graph_objects as go
warnings.filterwarnings('ignore')
from st_aggrid import AgGrid, GridUpdateMode, JsCode
from st_aggrid.grid_options_builder import GridOptionsBuilder
from st_aggrid import GridOptionsBuilder, AgGrid, GridUpdateMode, DataReturnMode
from io import BytesIO
from streamlit_folium import st_folium
#from pyxlsb import open_workbook as open_xlsb
import plotly.figure_factory as ff
import plotly.express as px
import folium
import streamlit.components.v1 as components
from pandas.api.types import (
    is_categorical_dtype,
    is_datetime64_any_dtype,
    is_numeric_dtype,
    is_object_dtype,
)
import math
import datetime
from streamlit_option_menu import option_menu
from sign import *
from base import *
from fpdf import FPDF
import base64
import docx
from pathlib import Path
from docx.shared import Inches
from docx.shared import RGBColor


#@st.experimental_memo
def load_data(text):
    df = pd.read_excel(text)
    return df

if "projet" not in st.session_state:
        st.session_state.projet = dict()
if "liste_projet" not in st.session_state:
        st.session_state.liste_projet = (load_data("liste_projects.xlsx")).to_dict(orient='records')  
if "c_cours" not in st.session_state:
        st.session_state.c_cours = False



def st_directory_picker(initial_path=Path()):

    st.markdown("Choisir un répertoire")

    if "path" not in st.session_state:
        st.session_state.path = initial_path.absolute()

    manual_input = st.text_input("Le dossier choisi:", st.session_state.path)

    manual_input = Path(manual_input)
    if manual_input != st.session_state.path:
        st.session_state.path = manual_input
        st.experimental_rerun()

    _, col1, col2, col3, _ = st.columns([3, 1, 3, 1, 3])

    with col1:
        st.markdown("#")
        if st.button("⬅️") and "path" in st.session_state:
            st.session_state.path = st.session_state.path.parent
            st.experimental_rerun()

    with col2:
        subdirectroies = [
            f.stem
            for f in st.session_state.path.iterdir()
            if f.is_dir()
            and (not f.stem.startswith(".") and not f.stem.startswith("__"))
        ]
        if subdirectroies:
            st.session_state.new_dir = st.selectbox(
                "Sous dossiers", sorted(subdirectroies)
            )
        else:
            st.markdown("#")
            st.markdown(
                "<font color='#FF0000'>No subdir</font>", unsafe_allow_html=True
            )

    with col3:
        if subdirectroies:
            st.markdown("#")
            if st.button("➡️") and "path" in st.session_state:
                st.session_state.path = Path(
                    st.session_state.path, st.session_state.new_dir
                )
                st.experimental_rerun()
    
    return str(st.session_state.path)
   

if "qte" not in st.session_state:
        myFile = open("quantite_projet.xlsx", "w+")
        dataframe=pd.DataFrame(columns=['Sous Système', 'N° préstation', 'Désignation','Travaux','Quantité','Taux forfaitaire unitaire JOUR',"Taux forfaitaire unitaire NUIT LONGUE","Fournitures unitaires","Temps de main d'oeuvre","CMP","Délai d'appro","SAV"])
        st.session_state.qte = dataframe
        dataframe.to_excel("quantite_projet.xlsx",index=False)

if "estimation1" not in st.session_state:
    st.session_state.estimation1 =pd.DataFrame()
 
if "estimation2" not in st.session_state:
    st.session_state.estimation2 =pd.DataFrame()




#@st.cache
def convert_df(df):
    # IMPORTANT: Cache the conversion to prevent computation on every rerun
    return df.to_csv(index=False).encode("utf-8")


def check_password(username,password):
    """Returns `True` if the user had a correct password."""

    def password_entered():
        """Checks whether a password entered by the user is correct."""
        if (
            st.session_state["username"] in st.secrets["passwords"]
            and st.session_state["password"]
            == st.secrets["passwords"][st.session_state["username"]]
        ):
            st.session_state["password_correct"] = True
            del st.session_state["password"]  # don't store username + password
            del st.session_state["username"]
        else:
            st.session_state["password_correct"] = False
    res=False
    if "password_correct" not in st.session_state:
        # First run, show inputs for username + password.
        st.text_input("Username", on_change=password_entered, key="username")
        st.text_input(
            "Password", type="password", on_change=password_entered, key="password"
        )
        if st.button('OK'):
             return res
    elif not st.session_state["password_correct"]:
        # Password not correct, show input + error.
        st.text_input("Username", on_change=password_entered, key="username")
        st.text_input(
            "Password", type="password", on_change=password_entered, key="password"
        )
        if st.button('OK'): 
            st.error("😕 User not known or password incorrect")
            return False
    else:
        res=True
        return res
    
def password_enteredV1(username,password):
        """Checks whether a password entered by the user is correct."""
        if (
            username in st.secrets["passwords"]
            and password
            == st.secrets["passwords"][username]
        ):
            return True
        else:
            return False

    

def show_grid(df):
    gb = GridOptionsBuilder.from_dataframe(df)
    gb.configure_default_column(editable=True)
    grid_table = AgGrid(
        df,
        height = "800px", 
        width='100%',
        gridOptions=gb.build(),
        fit_columns_on_grid_load=False,
        allow_unsafe_jscode=True,
    )
    return grid_table

def show_grid1(df):
    gb = GridOptionsBuilder.from_dataframe(df)
    gb.configure_default_column(editable=False)
    grid_table = AgGrid(
        df,
        height = "800px", 
        width='100%',
        gridOptions=gb.build(),
        fit_columns_on_grid_load=False,
        allow_unsafe_jscode=True,
    )
    return grid_table


def update(df):
    grid_table=show_grid(df)
    grid_table_df = pd.DataFrame(grid_table['data'])
    return grid_table_df


def f1() :   
            data=st.session_state.data
            res = option_menu("BPU", ['Consulter', 'Rechercher','Ajouter','Modifier','Supprimer','Extraire'],key="id6",
                                 icons=['house', 'list-task', 'server', 'pencil','trash','upload'],
                                 menu_icon="cast", default_index=0,orientation="horizontal",
                                 styles={
                "container": {"padding": "5!important", "background-color": "#5cb8a7"},
                "icon": {"color": "blue", "font-size": "20px"}, 
                "nav-link": {"font-size": "16px", "text-align": "left", "margin":"2px", "--hover-color": "#c8dfe3"},
                "nav-link-selected": {"background-color": "#034980"},
                      }
                )

            if (res=='Consulter'):
                tab1, tab2,tab3= st.tabs(["Préstations", "Consultation des Sous Systèmes","Mise a jour des Sous Systèmes"])
                dat=st.session_state.syst
                with tab1:
                        gb = GridOptionsBuilder.from_dataframe(data)
                        gb.configure_pagination(paginationAutoPageSize=True) #Add pagination
                        gb.configure_side_bar() #Add a sidebar
                        gb.configure_selection('disabled', use_checkbox=True, groupSelectsChildren="Group checkbox select children") 
                        gridOptions = gb.build()
                        grid_response = AgGrid(
                            data,
                            gridOptions=gridOptions,
                            data_return_mode='AS_INPUT', 
                            update_mode='MODEL_CHANGED', 
                            fit_columns_on_grid_load=False,
                            enable_enterprise_modules=True,
                            theme='alpine',
                            height = "800px", 
                            width='100%',
                            reload_data=False
                        )

                        data = grid_response['data']
                        selected = grid_response['selected_rows'] 
                        df = pd.DataFrame(selected)
                with tab2:
                        gb = GridOptionsBuilder.from_dataframe(dat)
                        gb.configure_pagination(paginationAutoPageSize=True) #Add pagination
                        gb.configure_side_bar() #Add a sidebar
                        gb.configure_selection('disabled', use_checkbox=True, groupSelectsChildren="Group checkbox select children") 
                        gridOptions = gb.build()
                        grid_response = AgGrid(
                            dat,
                            gridOptions=gridOptions,
                            data_return_mode='AS_INPUT', 
                            update_mode='MODEL_CHANGED', 
                            fit_columns_on_grid_load=False,
                            enable_enterprise_modules=True,
                            theme='alpine',
                            height = "800px", 
                            width='100%',
                            reload_data=False
                        )

                        data = grid_response['data']
                        selected = grid_response['selected_rows'] 
                        df = pd.DataFrame(selected)
                        df_xlsx = to_excell(dat)
                        st.download_button(label='📥 Télécharger',
                                                data=df_xlsx ,
                                                file_name= 'SOUS_SYSTEMES.xlsx')
                with tab3: 
                        dat = dat.astype(str)
                        res = st.radio("Choisir : ", ('Rechercher 🕵️‍♂️','Ajouter 👈'))
                        if res=='Rechercher 🕵️‍♂️': 
                            st.dataframe(filter_dataframe(dat))
                        else:
                            d=dict()
                            num=st.text_input("N°Sous Système :")
                            des=st.text_input("Désignation :")
                            d["N°Sous Système"]=str(num)
                            d["Désignation"]=des
                            df_dictionary = pd.DataFrame([d])
                            if st.button("Ajouter"):
                                s=st.session_state.syst
                                s = s.astype(str)
                                if str(num) not in (s["N°Sous Système"].unique()):
                                        dat = pd.concat([dat, df_dictionary], ignore_index=True)
                                        st.session_state.syst = dat
                                        dat.to_excel("Sous_Systeme.xlsx",index=False)
                                        st.success('Ajout éffectué avec succés!!!')
                                else:
                                    st.error('Numero de sous système déja existant!!!')

            elif (res=='Ajouter'):
                    if st.session_state.admin:
                           data=user_add_pres(data)
                           st.session_state.data = data
                           (st.session_state.data).to_excel('BPU'+str(st.session_state.titu)+'.xlsx',index=False)
                    else:
                         st.error("Vous n'etes pas admin")

            elif (res=='Rechercher'):
                   rech=filter_dataframe(data)
                   st.dataframe(rech)
                   df_xlsx = to_excell(rech)
                   st.download_button(label='📥 Télécharger',
                                                data=df_xlsx ,
                                                file_name= 'Prestations_filtrées.xlsx')
            elif (res=='Modifier'):
                if st.session_state.admin:
                     data=update(data)
                     st.session_state.data = data 
                     (st.session_state.data).to_excel('BPU'+str(st.session_state.titu)+'.xlsx',index=False)
                else:
                    st.error("Vous n'etes pas admin")
            elif (res=='Supprimer'):
                if st.session_state.admin:
                    data=user_supp_pres(data)
                    st.session_state.data = data 
                    (st.session_state.data).to_excel('BPU'+str(st.session_state.titu)+'.xlsx',index=False)
                else:
                    st.error("Vous n'etes pas admin")
            else:
                df_xlsx = to_excell(data)
                with st.container():
                    col1,col2,col3,col4,col5=st.columns(5)
                    with col3:
                          st.download_button(label='Extraire',data=df_xlsx ,file_name='BPU_'+str(st.session_state.titu)+'.xlsx')
    
            
def user_add_pres(data):
    with st.container():
        st.subheader("Ajouter une préstation")
        st.markdown(
    """
<style>
    div[data-testid="columns"] {
    box-shadow: rgb(0 0 0 / 20%) 0px 2px 1px -1px, rgb(0 0 0 / 14%) 0px 1px 1px 0px, rgb(0 0 0 / 12%) 0px 1px 3px 0px;
    border-radius: 15px;
    padding: 5% 5% 5% 10%;
} 
</style>
""",
    unsafe_allow_html=True,
)
        col1,col2 = st.columns(2)
        dd = st.session_state.syst
        d=dict()
        with col1:
                sys=st.selectbox('Sous système:',dd["Désignation"].unique())
                des=dd[dd["Désignation"]==sys]
                des=(des["N°Sous Système"].unique())[0]
                st.write('Sous système concerné :' ,str(des))
                liste=((data[data['Sous Système']==des])["Type préstation"]).unique()
                prestation=st.selectbox('Type de la préstation:',liste)
                num_prix= st.text_input("N°préstation :")
                designation= st.text_input("Désignation :")
                unite=st.text_input("Unité:",'u')
                
        with col2:
                fourniture= st.number_input('Fournitures(€):')
                mo= st.number_input("Temps de main d'oeuvre (heures):")
                mo_jour_h= st.number_input("Prix unitaire MO JOUR |Taux horaire (€):")
                mo_nuit_ch= st.number_input("Prix unitaire MO NUIT COURTE |Taux horaire (€):")
                mo_nuit_lh= st.number_input("Prix unitaire MO NUIT LONGUE |Taux horaire (€):")
        d["N° de prix "]=num_prix
        d["Désignation"]=designation
        d["Unité"]=unite
        d["Sous Système"]=des
        d["Type préstation"]=prestation
        d['Fournitures \nP.U en euros']=fourniture
        d["Temps Main d'œuvre en heures"]=mo
        d['Prix unitaire MO JOUR (hors fourniture)|Taux horaire']=mo_jour_h
        v1=float(mo)*float(mo_jour_h)
        d['Prix unitaire MO JOUR (hors fourniture)|Taux forfaitaire']=float(mo)*float(mo_jour_h)
        d['Prix unitaire MO NUIT COURTE (hors fourniture)|Taux horaire']=mo_nuit_ch
        v2=float(mo)*float(mo_nuit_ch)
        d['Prix unitaire MO NUIT COURTE (hors fourniture)|Taux forfaitaire']=float(mo)*float(mo_nuit_ch)
        d['Prix unitaire MO NUIT LONGUE (hors fourniture)|Taux horaire']=mo_nuit_lh
        v3=float(mo)*float(mo_nuit_lh)
        d['Prix unitaire MO NUIT LONGUE (hors fourniture)|Taux forfaitaire']=float(mo)*float(mo_nuit_lh)
        if st.button('Ajouter ✅'):
            t=data.astype(str)
            if str(num_prix) not in (t["N° de prix "].unique()):
                st.success('Ajout éffectué avec succés!!!')
                df_dictionary = pd.DataFrame([d])
                data = pd.concat([data, df_dictionary], ignore_index=True)
                data.reset_index(drop=True, inplace=True)
                st.write(data)
                return data  
            else:
                st.error('Numéro de préstation déja éxistant!!!')
                return data
        else:
            return data
    
            


def f2() :   
    data=st.session_state.eq                                         
    res = option_menu("EQUIPEMENTS ATLAS", ['Consulter', 'Rechercher','Ajouter','Modifier','Supprimer','Extraire'],key="id5",
                                 icons=['house', 'list-task', 'server', 'pencil','trash','upload'],
                                 menu_icon="cast", default_index=0,orientation="horizontal",
                                 styles={
                "container": {"padding": "5!important", "background-color": "#5cb8a7"},
                "icon": {"color": "blue", "font-size": "20px"}, 
                "nav-link": {"font-size": "16px", "text-align": "left", "margin":"0px", "--hover-color": "#c8dfe3"},
                "nav-link-selected": {"background-color": "#034980"},
                      }
                )
    if (res=='Consulter'):
        gb = GridOptionsBuilder.from_dataframe(data)
        gb.configure_pagination(paginationAutoPageSize=True) #Add pagination
        gb.configure_side_bar() #Add a sidebar
        gb.configure_selection('disabled', use_checkbox=True, groupSelectsChildren="Group checkbox select children") #Enable multi-row selection
        gridOptions = gb.build()
        grid_response = AgGrid(
            data,
            gridOptions=gridOptions,
            data_return_mode='AS_INPUT', 
            update_mode='MODEL_CHANGED', 
            fit_columns_on_grid_load=False,
            theme='alpine', #Add theme color to the table
            enable_enterprise_modules=True,
            height = "800px", 
            width='100%',
            reload_data=False
        )

        data = grid_response['data']
        selected = grid_response['selected_rows'] 
        df = pd.DataFrame(selected)
    elif (res=='Ajouter'):
        if st.session_state.admin:
               st.write('''<style>
    [data-testid="stHorizontalBlock"]:has(div.PortMarker) [data-testid="stMarkdownContainer"] p { 
        margin: 0px 0px 0.2rem; 
        color: #ff0000;
    }        
    </style>''', unsafe_allow_html=True)
               with st.container():
                   data=user_add_eq(data)
                   st.session_state.eq = data
                   (st.session_state.eq).to_excel('Equipements.xlsx',index=False)
        else:
                    st.error("Vous n'etes pas admin")
             
    elif (res=='Rechercher'): 
           rech=filter_dataframe(data)
           st.dataframe(rech)
           df_xlsx = to_excell(rech)
           st.download_button(label='📥 Télécharger',
                                                data=df_xlsx ,
                                                file_name= 'Equipements_filtrées.xlsx')
    elif (res=='Modifier'):
        if st.session_state.admin:
             data=update(data)
             st.session_state.eq = data 
             (st.session_state.eq).to_excel('Equipements.xlsx',index=False)
        else:
                    st.error("Vous n'etes pas admin")
    elif (res=='Supprimer') :
        if st.session_state.admin:
            data=user_supp_eq(data)
            st.session_state.eq = data 
            (st.session_state.eq).to_excel('Equipements.xlsx',index=False)
        else:
                    st.error("Vous n'etes pas admin")
    else:
        df_xlsx = to_excell(st.session_state.eq)
        st.write('''<style>
    [data-testid="stHorizontalBlock"]:has(div.PortMarker) [data-testid="stMarkdownContainer"] p { 
        margin: 0px 0px 0.2rem; 
        color: #ff0000;
    }        
    </style>''', unsafe_allow_html=True)
        with st.container():
            col1,col2,col3,col4,col5=st.columns(5)
            with col3:
                st.download_button(label='Extraire',data=df_xlsx ,file_name= 'Equipements.xlsx')
        
        
        

def table_interactive(text):
        data = load_data(text)
        gb = GridOptionsBuilder.from_dataframe(data)
        gb.configure_pagination(paginationAutoPageSize=True) #Add pagination
        gb.configure_side_bar() #Add a sidebar
        gb.configure_selection('multiple', use_checkbox=True, groupSelectsChildren="Group checkbox select children") #Enable multi-row selection
        gridOptions = gb.build()
        grid_response = AgGrid(
            data,
            gridOptions=gridOptions,
            data_return_mode='AS_INPUT', 
            update_mode='MODEL_CHANGED', 
            fit_columns_on_grid_load=False,
            theme='alpine', #Add theme color to the table
            enable_enterprise_modules=True,
            height = "800px", 
            width='100%',
            reload_data=False
        )

        data = grid_response['data']
        selected = grid_response['selected_rows'] 
        df = pd.DataFrame(selected)
        return df
def filter_dataframe(df: pd.DataFrame) -> pd.DataFrame:
    modify = st.checkbox("Add filters")
    if not modify:
        return df
    df = df.copy()
    for col in df.columns:
        if is_object_dtype(df[col]):
            try:
                df[col] = pd.to_datetime(df[col])
            except Exception:
                pass

        if is_datetime64_any_dtype(df[col]):
            df[col] = df[col].dt.tz_localize(None)

    modification_container = st.container()

    with modification_container:
        to_filter_columns = st.multiselect("Filter dataframe on", df.columns)
        for column in to_filter_columns:
            left, right = st.columns((1, 20))
            # Treat columns with < 10 unique values as categorical
            if is_categorical_dtype(df[column]) or df[column].nunique() < 10:
                user_cat_input = right.multiselect(
                    f"Values for {column}",
                    df[column].unique(),
                    default=list(df[column].unique()),
                )
                df = df[df[column].isin(user_cat_input)]
            elif is_numeric_dtype(df[column]):
                _min = float(df[column].min())
                _max = float(df[column].max())
                step = (_max - _min) / 100
                user_num_input = right.slider(
                    f"Values for {column}",
                    min_value=_min,
                    max_value=_max,
                    value=(_min, _max),
                    step=step,
                )
                df = df[df[column].between(*user_num_input)]
            elif is_datetime64_any_dtype(df[column]):
                user_date_input = right.date_input(
                    f"Values for {column}",
                    value=(
                        df[column].min(),
                        df[column].max(),
                    ),
                )
                if len(user_date_input) == 2:
                    user_date_input = tuple(map(pd.to_datetime, user_date_input))
                    start_date, end_date = user_date_input
                    df = df.loc[df[column].between(start_date, end_date)]
            else:
                user_text_input = right.text_input(
                    f"Substring or regex in {column}",
                )
                if user_text_input:
                    df = df[df[column].astype(str).str.contains(user_text_input)]

    return df
        
        
        
        
        

            

def user_add_eq(data):
        st.subheader("Ajouter un équipement")
        col1,col2 = st.columns(2)
        d=dict()
        with col1:
                ref=st.text_input("Référence Article:")
                d["Référence Article"]=ref
                designation= st.text_input("Libellé Article :")
                d["Libellé Article"]=designation
                catalogue= st.text_input("Catalogue :")
                d["Catalogue"]=catalogue
                famille= st.text_input("Famille :")
                d["Famille"]=famille
                ssfamille= st.text_input("Sous Famille :")
                d["Sous-Famille"]=ssfamille
                usage= st.text_input("Usage :")
                d["Usage"]=usage
                delai= st.number_input("Délai d'approvisionnement (jours):")
                d["Délai d'appro (jours)"]=delai
                cmp= st.number_input("CMP (€):")
                d["CMP (€)"]=cmp
        with col2:
                
                fournisseur= st.text_input("Fournisseur :")
                d["Fournisseur"]=fournisseur
                marche= st.text_input("N° de marché :")
                d["N° de marché"]=marche
                fabricant= st.text_input("Fabricant :")
                d["Fabricant"]=fabricant
                sav= st.number_input("SAV (€):")
                d["SAV"]=sav
                comment= st.text_area("Commentaire achat:")
                d["libelleAchat"]=comment
                dd =st.session_state.syst
                sys=st.selectbox('Sous système:',dd['Désignation'].unique())
                des=dd[dd["Désignation"]==sys]
                des=(des["N°Sous Système"].unique())[0]
                st.write('Sous système concerné :' ,str(des))
                d["Sous Système"]=des
        if st.button('Ajouter ✅'):
            res =data.astype(str)
            if str(ref) not in res["Référence Article"].unique():
                st.success('Ajout éffectué avec succés!!!')
                df_dictionary = pd.DataFrame([d])
                data = pd.concat([data, df_dictionary], ignore_index=True)
                data.reset_index(drop=True, inplace=True)
                st.write(data)
                return data  
            else:
                st.error('Référence article déja éxistante!!!')
                return data
        else:
            return data

        
def user_supp_pres(data):
        st.subheader("Supprimer des préstations")
        gb = GridOptionsBuilder.from_dataframe(data)
        gb.configure_pagination(paginationAutoPageSize=True) #Add pagination
        gb.configure_side_bar() #Add a sidebar
        gb.configure_selection('multiple', use_checkbox=True, groupSelectsChildren="Group checkbox select children") #Enable multi-row selection
        gridOptions = gb.build()
        grid_response = AgGrid(
            data,
            gridOptions=gridOptions,
            data_return_mode='AS_INPUT', 
            update_mode='MODEL_CHANGED', 
            fit_columns_on_grid_load=False,
            theme='alpine', #Add theme color to the table
            enable_enterprise_modules=True,
            height = "800px", 
            width='100%',
            reload_data=False
        )

        data = grid_response['data']
        selected = grid_response['selected_rows'] 
        df_selected = pd.DataFrame(selected)
        if st.button('Supprimer ❌'):
                df_selected=df_selected.set_index("N° de prix ")
                for elem in list(df_selected.index):
                        data=data[data["N° de prix "]!=elem]
                st.success('Suppression éffectuée avec succés!!!')
        else:
            pass
            
        return data
def user_supp_eq(data):
        st.subheader("Supprimer des équipements")
        gb = GridOptionsBuilder.from_dataframe(data)
        gb.configure_pagination(paginationAutoPageSize=True) #Add pagination
        gb.configure_side_bar() #Add a sidebar
        gb.configure_selection('multiple', use_checkbox=True, groupSelectsChildren="Group checkbox select children") #Enable multi-row selection
        gridOptions = gb.build()
        grid_response = AgGrid(
            data,
            gridOptions=gridOptions,
            data_return_mode='AS_INPUT', 
            update_mode='MODEL_CHANGED', 
            fit_columns_on_grid_load=False,
            theme='alpine', #Add theme color to the table
            enable_enterprise_modules=True,
            height = "800px", 
            width='100%',
            reload_data=False
        )

        data = grid_response['data']
        selected = grid_response['selected_rows'] 
        df_selected = pd.DataFrame(selected)
        if st.button('Supprimer ❌'):
                df_selected=df_selected.set_index("Référence Article")
                for elem in list(df_selected.index):
                        data=data[data["Référence Article"]!=elem]
                st.success('Suppression éffectuée avec succés!!!')
        else:
            pass
            
        return data
def user_supp_qte():
        data=st.session_state.qte
        st.subheader("Supprimer des quantités")
        gb = GridOptionsBuilder.from_dataframe(data)
        gb.configure_pagination(paginationAutoPageSize=True) #Add pagination
        gb.configure_side_bar() #Add a sidebar
        gb.configure_selection('multiple', use_checkbox=True, groupSelectsChildren="Group checkbox select children")
        gridOptions = gb.build()
        grid_response = AgGrid(
                data,
                gridOptions=gridOptions,
                data_return_mode='AS_INPUT', 
                update_mode='MODEL_CHANGED', 
                fit_columns_on_grid_load=False,
                theme='alpine',
                enable_enterprise_modules=True,
                height = "800px", 
                width='100%',
                reload_data=False
                )
        data= grid_response['data']
        if st.button('Supprimer ❌'):
                selected = grid_response['selected_rows']        
                df_selected = pd.DataFrame(selected)
               
                data=data.drop(df_selected.index)
               
                st.success('Suppression éffectuée avec succés!!!')
                st.write(data)
        else:
            pass
        return data
def association():
                data=st.session_state.data
                eq=st.session_state.eq
                res = option_menu("EQUIPEMENTS ATLAS", ['Consulter', 'Rechercher','Ajouter','Modifier','Supprimer'],key="id33",
                                 icons=['house', 'list-task', 'server', 'pencil','trash'],
                                 menu_icon="cast", default_index=0,orientation="horizontal",
                                 styles={
                "container": {"padding": "5!important", "background-color": "#5cb8a7"},
                "icon": {"color": "blue", "font-size": "20px"}, 
                "nav-link": {"font-size": "16px", "text-align": "left", "margin":"0px", "--hover-color": "#c8dfe3"},
                "nav-link-selected": {"background-color": "#034980"},
                      }
                )
                if res=='Consulter':
                    mdata=load_data("prestation_equipement.xlsx")
                    if (mdata.shape[0] > 0):
                            
                            d1=st.session_state.data
                            d2=st.session_state.eq
                            ll=[]
                            for i in range(len(mdata)):
                                    d=dict()
                                    d["N° prix"]=mdata["N° de prix "][i]
                                    d["Préstation"]=(((d1[d1["N° de prix "]==d["N° prix"]])["Désignation"]).unique())[0]
                                    d["Référence Article"]=mdata["Référence Article"][i]
                                    d["Equipement"]=(((d2[d2["Référence Article"]==d["Référence Article"]])["Libellé Article"]).unique())[0]
                                    ll.append(d)
                            ll=pd.DataFrame(ll)
                            st.dataframe(filter_dataframe(ll)) 
                            df_xlsx = to_excell(ll)
                            st.download_button(label='📥 Télécharger',
                                        data=df_xlsx ,
                                        file_name= 'PRESTATION-EQUIPEMENT.xlsx')
                            agree = st.checkbox('Filtrage par préstation',key='teest')
                            if agree:
                                prestation=st.selectbox('Préstation:',ll["Préstation"].unique())
                                if st.button("OK"):
                                    st.dataframe(ll[ll["Préstation"]==prestation])
                            
                    else:
                        st.warning('Aucune association trouvée!!!')
                        
                            
                        
                         
                elif res=='Rechercher':
                    
                    mdata=st.session_state.soc
                    if (mdata.shape[0] > 0):
                        st.dataframe(filter_dataframe(st.session_state.soc))
                    else:
                        st.warning('Aucune association trouvée!!!')
                elif res =='Ajouter':
                    if st.session_state.admin:
                        dd = st.session_state.syst
                        sys=st.selectbox('Sous systeme:',dd["Désignation"].unique())
                        des=dd[dd["Désignation"]==sys]
                        des=(des["N°Sous Système"].unique())[0]
                        st.write('Sous système concerné :' ,str(des))
                        liste=((data[data['Sous Système']==des])["Désignation"]).unique()
                        prestation=st.selectbox('Prestation:',liste)
                        ll=(data[data['Désignation']==prestation])
                        ll=(ll["N° de prix "].unique())[0]
                        eqq=eq[eq['Sous Système']==des]
                        gb = GridOptionsBuilder.from_dataframe(eqq)
                        gb.configure_pagination(paginationAutoPageSize=True) #Add pagination
                        gb.configure_side_bar() #Add a sidebar
                        gb.configure_selection('multiple', use_checkbox=True, groupSelectsChildren="Group checkbox select children") 
                        gridOptions = gb.build()
                        grid_response = AgGrid(
                        eqq,
                        gridOptions=gridOptions,
                        data_return_mode='AS_INPUT', 
                        update_mode='MODEL_CHANGED', 
                        fit_columns_on_grid_load=False,
                        theme='alpine', #Add theme color to the table
                        enable_enterprise_modules=True,
                        height = "800px", 
                        width='100%',
                        reload_data=False
                )

                        eqq = grid_response['data']
                        selected = grid_response['selected_rows'] 
                        df = pd.DataFrame(selected)
                        dd = pd.read_excel("prestation_equipement.xlsx")
                        if st.button("Associer ✅"):
                            l1=[]
                            l2=[]
                            if (df.shape[0]) >0 :
                                    for elem in df['Référence Article'] :
                                        if (ll,elem) not in zip(dd["N° de prix "],dd['Référence Article']):
                                          l1.append(elem)
                                          l2.append(ll)
                                        else:
                                                st.warning('Association '+str(ll)+ " - "+str(elem)+ " déja éxistante!!!")
                                    zipped = list(zip(l2, l1))
                                    df = pd.DataFrame(zipped, columns=["N° de prix ", 'Référence Article'])
                                    dd = pd.concat([dd, df], ignore_index=True)
                                    dd.to_excel('prestation_equipement.xlsx',index=False)
                                    st.success('Association éffectuée avec succés!!!')
                                    st.write(dd)
                                    st.session_state.soc=dd
                                    df_xlsx = to_excell(dd)
                                    st.download_button(label='📥 Télécharger',
                                        data=df_xlsx ,
                                        file_name= 'PRESTATION_EQUIPEMETS.xlsx')
                            else:
                                    st.warning("Aucun équipement a associer")
                    else:
                         st.error("Vous n'etes pas admin")
                else :
                    if st.session_state.admin:
                        res=(st.session_state.soc).copy()
                        liste=res["N° de prix "].unique()
                        num=st.selectbox('Référence préstation:',liste)
                        if st.button("Supprimer cette association"):
                                st.session_state.soc=res[(res["N° de prix "]!=num) ]
                                st.success('Association supprimée avec succés!!!')
                                (st.session_state.soc).to_excel('prestation_equipement.xlsx',index=False)
                    else:
                          st.error("Vous n'etes pas admin")
                            
                            
def check_project(my_list,num):
    for dico in my_list:
        if dico['N° Projet'] ==num:
            return dico
    return dict()

def manage_projects():
        
        res= option_menu("Chiffrage", ['Nouveau','En cours','Précédent'],key="id1073",
                                                         icons=['house', 'list-task'],
                                                         menu_icon="app-indicator", default_index=0,orientation='horizontal',
                                                         styles={
                                        "container": {"padding": "5!important", "background-color": "#5cb8a7"},
                                        "icon": {"color": "blue", "font-size": "25px"}, 
                                        "nav-link": {"font-size": "16px", "text-align": "left", "margin":"0px", "--hover-color": "#c8dfe3"},
                                        "nav-link-selected": {"background-color": "#034980"},
                                              }
                                        )
                                        
        if res=='En cours' :
            b=st.session_state.projet
            if (st.session_state.c_cours)==True:
                    if b!=dict():
                        st.metric('Nom du projet',b['Nom du Projet'])
                    tit=st.session_state.titu
                    if tit=='EIFFAGE':
                           st.session_state.data = load_data("BPU_EIFFAGE.xlsx")
                    elif tit =='ERI':
                          st.session_state.data = load_data("BPU_ERI.xlsx")
                    elif tit =='INEO':
                         st.session_state.data = load_data("BPU_INEO.xlsx")
                    elif tit =='SDEL':
                         st.session_state.data = load_data("BPU_SDEL.xlsx") 
                    else:
                         st.session_state.data = load_data("BPU_SOGETREL.xlsx")
                    
                    manage_quantite()
            else:
                 st.error('Aucun chiffrage en cours')
        if res =='Précédent':
                        d=[]
                        dataframe=pd.DataFrame()
                        folder_path = "./historique"
                        for path, dirs, files in os.walk(folder_path):
                             for filename in files:
                                    if filename.split('.')[1] =='xlsx':
                                        d.append(filename)
                        if len(d)==0:
                                st.error("Historique vide")
                        else:
                                 file=st.selectbox('Sélectionner:',d)
                                 if st.button('Valider'):
                                        st.session_state.qte=load_data(folder_path+'/'+file)
                                        num_projet=file.split('_')[0]
                                        rest=check_project(st.session_state.liste_projet,num_projet)
                                        if rest !=dict():
                                            st.session_state.projet=rest
                                            (st.session_state.qte).to_excel('quantite_projet.xlsx',index=False)
                                        st.success('Le chiffrage a été restauré')
                                        st.session_state.c_cours=True
                                        
        if res=='Nouveau':
                           nouveau_chiffrage()
                           
                           
                            
                           
                           
                            
def nouveau_chiffrage():
                
                ddf=pd.read_csv('gares.csv',header=0,sep=';')
                ddf=(ddf[ddf['exploitant']=='RATP'])
                dictio=dict()
                col1,col2=st.columns(2)
                with col1:
                        nom_projet= st.text_input("Nom du projet :",key='00')
                        num_projet= st.text_input("N° Projet :",key='01')
                        reseau=st.selectbox('Réseau :',['METRO','RER','BUS','TRAMWAY','BATIMENT','ATELIER'])
                        if reseau =='RER':
                                ligne=st.selectbox('Ligne :',['RER A',' RER B','RER C','RER D'])
                        elif reseau=='METRO':
                                ligne=st.selectbox('Ligne :',['M01','M02','M03','M04','M05','M06','M07','M08','M09','M10','M11','M12','M13','M14'])
                        else:
                             ligne=st.selectbox('Ligne :',[])
                        phase=st.selectbox('Phase du projet:',['FAI','AVP','PRO','DER','OLS'])
                        d = st.date_input("Date:",datetime.date.today(),key='02')
                        version=st.text_input("Version :",key='03')
                with col2:
                        ddf=(ddf[ddf['mode_']==reseau])
                        lieu=st.selectbox('Lieu:',ddf['nom_long'].unique())
                        with st.container():
                            coord=((ddf[ddf['nom_long']==lieu])['Geo Point'].unique())[0]
                            lignee=((ddf[ddf['nom_long']==lieu])['res_com'].unique())[0]
                            coord=coord.split(',')
                            m = folium.Map(location=coord, zoom_start=10)
                            popup = folium.Popup(f"""  Lieu:{lieu}<br> <br>Ligne: {lignee}<br>""",max_width=250,)               
                            folium.Marker(coord, popup=popup).add_to(m)
                            st_data = st_folium(m, width=350,height=450)
                dictio['Nom du Projet']=nom_projet
                dictio['N° Projet']=num_projet
                dictio['Phase du projet']=phase
                dictio['Réseau']=reseau
                dictio['Ligne']=ligne
                dictio['Lieu']=lieu
                dictio['Date']=str(d)
                dictio['Version']=version
                col1,col2,col3=st.columns(3)
                with col2:
                                if st.button('Valider'):
                                        
                                        myFile = open("quantite_projet.xlsx", "w+")
                                        dataframe=pd.DataFrame(columns=['Sous Système', 'N° préstation', 'Désignation','Travaux','Quantité','Taux forfaitaire unitaire JOUR',"Taux forfaitaire unitaire NUIT LONGUE","Fournitures unitaires","CMP","Délai d'appro","SAV"])
                                        st.session_state.qte = dataframe
                                        dataframe.to_excel("quantite_projet.xlsx",index=False)
                                        st.success('Un nouveau chiffrage a été créé')
                                        st.session_state.c_cours=True
                                        st.session_state.projet=dictio
                                        
                                        
                                 
                                        
               
def manage_quantite():
        res = option_menu("QUANTITE DU PROJET", ['Consulter', 'Rechercher','Ajouter','Modifier','Supprimer','Extraire','Sauvegarder'],key="id5",
                                icons=['house', 'list-task', 'server', 'pencil','trash','upload','upload'],
                                 menu_icon="cast", default_index=0,orientation="horizontal",
                                 styles={
                "container": {"padding": "5!important", "background-color": "#5cb8a7"},
                "icon": {"color": "blue", "font-size": "15px"}, 
                "nav-link": {"font-size": "16px", "text-align": "left", "margin":"0px", "--hover-color": "#c8dfe3"},
                "nav-link-selected": {"background-color": "#034980"},
                      }
                )
        fusion=st.session_state.qte
        if res =='Consulter':
            if fusion.shape[0] >0:
                        gb = GridOptionsBuilder.from_dataframe(fusion)
                        gb.configure_pagination(paginationAutoPageSize=True) #Add pagination
                        gb.configure_side_bar() #Add a sidebar
                        gb.configure_selection('disabled', use_checkbox=True, groupSelectsChildren="Group checkbox select children") 
                        gridOptions = gb.build()
                        grid_response = AgGrid(
                            fusion,
                            gridOptions=gridOptions,
                            data_return_mode='AS_INPUT', 
                            update_mode='MODEL_CHANGED', 
                            fit_columns_on_grid_load=False,
                            enable_enterprise_modules=True,
                            theme='alpine',
                            height = "800px", 
                            width='100%',
                            reload_data=False
                        )

                        data = grid_response['data']
                        selected = grid_response['selected_rows'] 
                        df = pd.DataFrame(selected)
                        
                  
            else:
                st.error("Aucune quantité trouvée")
                    
        elif res=='Rechercher':
            if fusion.shape[0] >0:
                 
                 rech=filter_dataframe(fusion)
                 st.dataframe(rech)
                 df_xlsx = to_excell(rech)
                 st.download_button(label='Télécharger',
                                                data=df_xlsx ,
                                                file_name= 'Quantités_filtrées.xlsx')
            else:
                st.error("Aucune quantité trouvée")
        elif res=='Ajouter':
            quantite(fusion)
        elif res=='Modifier':
            if fusion.shape[0] >0:
                fusion=update(fusion)
                st.session_state.qte = fusion
                (st.session_state.qte).to_excel('quantite_projet.xlsx',index=False)
            else:
                st.error("Aucune quantité trouvée")
        elif res=='Supprimer':
                if fusion.shape[0] >0:
                    d=user_supp_qte()
                    st.session_state.qte=d
                    (d).to_excel('quantite_projet.xlsx',index=False)
                else:
                        st.error("Aucune quantité trouvée!!")
        elif res=='Extraire':
            if fusion.shape[0] >0:
                    df_xlsx = to_excell(fusion)
                    st.download_button(label='Extraire',data=df_xlsx ,file_name= 'QUANTITE PROJET.xlsx')
                    #st.succes("Extraction éffectuée avec succés")
            else:
                st.error("Aucune quantité trouvée")
        else:
            if fusion.shape[0] >0:
                  path=str(st.session_state.projet['Nom du Projet'])
                  if path =='':
                      path='Chiffrage_'+str(st.session_state.titu)
                  else:
                      path=str(st.session_state.projet['N° Projet'])+'_'+str(st.session_state.projet['Nom du Projet'])+'_'+str(st.session_state.titu)
                  sauvegarde(fusion,path+'.xlsx')
            else:
                st.error("Aucune quantité trouvée")
                       
def quantite(fusion):
                
                dictionnaire=dict()
                data=st.session_state.data
                eq=st.session_state.eq
                dd = st.session_state.syst
                qte=st.session_state.qte
                sys=st.selectbox('Sous système:',dd["Désignation"].unique())
                des=dd[dd["Désignation"]==sys]
                des=(des["N°Sous Système"].unique())[0]
                #st.write('Sous système concerné :' ,str(des))
                liste=((data[data['Sous Système']==des])["Désignation"]).unique()
                prestation=st.selectbox('Préstation:',liste)
                travaux=st.radio('Conditions de réalisation:',('NUIT COURTE','NUIT LONGUE','JOUR'),horizontal=True)
                qt= st.number_input("Quantité:",min_value=0)
                ll=(data[data['Désignation']==prestation])
                num_prestation=(ll["N° de prix "].unique())[0]
                dictionnaire["Sous Système"]=des
                dictionnaire["N° préstation"]=num_prestation
                dictionnaire["Désignation"]=prestation
                dictionnaire["Travaux"]=travaux
                dictionnaire["Quantité"]=qt
                dictionnaire["Taux forfaitaire unitaire JOUR"]=(ll["Prix unitaire MO JOUR (hors fourniture)|Taux forfaitaire"].unique())[0]
                dictionnaire["Taux forfaitaire unitaire NUIT COURTE"]=(ll["Prix unitaire MO NUIT COURTE (hors fourniture)|Taux forfaitaire"].unique())[0]
                dictionnaire["Taux forfaitaire unitaire NUIT LONGUE"]=(ll["Prix unitaire MO NUIT LONGUE (hors fourniture)|Taux forfaitaire"].unique())[0]
                dictionnaire["Fournitures unitaires"]=(ll["Fournitures \nP.U en euros"].unique())[0]
                dictionnaire["Temps de main d'oeuvre"]=(ll["Temps Main d'œuvre en heures"].unique())[0]
                dictionnaire= pd.DataFrame([dictionnaire])
                eqq=eq[eq['Sous Système']==des]                        
                gb = GridOptionsBuilder.from_dataframe(eqq)
                gb.configure_pagination(paginationAutoPageSize=True) #Add pagination
                gb.configure_side_bar() #Add a sidebar
                gb.configure_selection('multiple', use_checkbox=True, groupSelectsChildren="Group checkbox select children")
                gridOptions = gb.build()
                grid_response = AgGrid(
                eqq,
                gridOptions=gridOptions,
                data_return_mode='AS_INPUT', 
                update_mode='MODEL_CHANGED', 
                fit_columns_on_grid_load=False,
                theme='alpine',
                enable_enterprise_modules=True,
                height = "800px", 
                width='100%',
                reload_data=False
                )
                eqq = grid_response['data']
                selected = grid_response['selected_rows'] 
                df = pd.DataFrame(selected)
                dictionnaire["CMP"]=0
                dictionnaire["Délai d'appro"]=0
                dictionnaire["SAV"]=0
                if st.button("Valider ✅"):
                    s=df.shape
                    if (s[0] >0):
                        df =df[['Référence Article','Libellé Article','Catalogue','Famille',"Délai d'appro (jours)",'CMP (€)','Fournisseur','N° de marché','Fabricant','libelleAchat','Sous Système','SAV']]
                        df['Quantité']=qt
                        df['Num prestation']=num_prestation
                        df['Désignation prestation']=prestation
                        st.session_state.eq_asoc=load_data('Eq_asoc.xlsx')
                        fus=st.session_state.eq_asoc
                        if df.shape[0] >0:
                             fus= pd.concat([fus,df], ignore_index=True)
                             fus.to_excel('Eq_asoc.xlsx',index=False)   
                             st.session_state.eq_asoc=fus
                        dictionnaire["CMP"]=(df["CMP (€)"]).sum()
                        dictionnaire["SAV"]=(df["SAV"]).sum()
                        dictionnaire["Délai d'appro"]=(df["Délai d'appro (jours)"]).max()
                    fusion= pd.concat([fusion,dictionnaire], ignore_index=True)
                    fusion.to_excel('quantite_projet.xlsx',index=False)
                    st.success('Quantité ajoutée avec succés!!!')
                    st.session_state.qte=fusion
                    fusion= fusion.astype(str)
                    st.write(fusion)
                    
def to_excell(df):
    output = BytesIO()
    writer = pd.ExcelWriter(output, engine='xlsxwriter')
    df.to_excel(writer, index=False, sheet_name='Sheet1')
    workbook = writer.book
    worksheet = writer.sheets['Sheet1']
    format1 = workbook.add_format({'num_format': '0.00'}) 
    worksheet.set_column('A:A', None, format1)  
    writer.close()
    processed_data = output.getvalue()
    return processed_data

def get_dataframe(res):
                                my_list=res["N°Sous Système"].unique()
                                dd = st.session_state.syst
                                ll=[]
                                for elem in my_list:
                                    d=dict()
                                    my_data=res[res["N°Sous Système"]==elem]
                                    d["N°Sous Système"]=elem
                                    des=dd[dd["N°Sous Système"]==elem]
                                    des=(des["Désignation"].unique())[0]
                                    d["Désignation"]=des
                                    d["COUT FOURNITURE RATP"]=round(my_data["COUT FOURNITURE RATP"].sum(), 2)
                                    d["COUT FOURNITURE TITULAIRE"]=round(my_data["COUT FOURNITURE TITULAIRE"].sum(), 2)
                                    d["COUT MO"]=round(my_data["COUT MO"].sum(),2)
                                    d["DUREE DES TRAVAUX"]=round(my_data["Durée des travaux"].sum(),2)
                                    d["Délai d'appro"]=round(my_data["Délai d'appro"].max(),2)
                                    d["SAV"]=round(my_data["SAV"].sum(),2)
                                    d["COUT TOTAL"]=round(my_data["COUT TOTAL"].sum())
                                    ll.append(d)
                                dataframe=pd.DataFrame(ll)
                                return dataframe
def round_up(n, decimals = 0):
    multiplier = 10 ** decimals
    return math.ceil(n * multiplier) / multiplier
def rapport_synthèse(param):
    my_list=['EIFFAGE','ERI','INEO','SDEL','SOGETREL']
    my_dataframe=[]
    for elem in my_list:
        ddictio=dict()
        st.session_state.data = load_data("BPU_"+str(elem)+".xlsx")
        ddictio['Titulaire']=elem
        df=pd.read_excel("quantite_projet.xlsx")
        data=st.session_state.data
        ll=[]
        s=df.shape
        if (s[0] >0):
                    somme =0
                    somme1=0
                    somme2=0
                    for i in range(len(df)):
                            d=dict()
                            d["N°Sous Système"]=df["Sous Système"][i]
                            d["Désignation"]=df["Désignation"][i]
                            d["Travaux"]=df["Travaux"][i]
                            #d["Travaux"]=param
                            d["Quantité"]=df["Quantité"][i]
                            d["CMP préstation unitaire"]=df["CMP"][i]
                            d["Fournitures préstation unitaire"]=df["Fournitures unitaires"][i]
                            lll=(data[data['Désignation']==d["Désignation"]])
                            #print((lll["Prix unitaire MO NUIT COURTE (hors fourniture)|Taux forfaitaire"].unique())[0])
                            #st.write(df)
                            df["Taux forfaitaire unitaire JOUR"]=(lll["Prix unitaire MO JOUR (hors fourniture)|Taux forfaitaire"].unique())[0]
                            df["Taux forfaitaire unitaire NUIT COURTE"]=(lll["Prix unitaire MO NUIT COURTE (hors fourniture)|Taux forfaitaire"].unique())[0]
                            df["Taux forfaitaire unitaire NUIT LONGUE"]=(lll["Prix unitaire MO NUIT LONGUE (hors fourniture)|Taux forfaitaire"].unique())[0]
                            if param !='PAR DEFAUT':
                                 df["Travaux"][i]=param
                            if (df["Travaux"][i])== 'JOUR' :
                                d["COUT MO"]=round((int(df["Quantité"][i]))*(float(df["Taux forfaitaire unitaire JOUR"][i])),2)
                            elif (df["Travaux"][i])== 'NUIT COURTE':
                                d["COUT MO"]=round((int(df["Quantité"][i]))*(float(df["Taux forfaitaire unitaire NUIT COURTE"][i])),2)
                            else:
                                d["COUT MO"]=round((int(df["Quantité"][i]))*(float(df["Taux forfaitaire unitaire NUIT LONGUE"][i])),2)
                            d["COUT FOURNITURE RATP"] =round(int(df["Quantité"][i])*float(df["CMP"][i]),2)
                            d["COUT FOURNITURE TITULAIRE"] =round(int(df["Quantité"][i])*float(df["Fournitures unitaires"][i]),2)
                            d["Durée des travaux"] =round(int(df["Quantité"][i])*float(df["Temps de main d'oeuvre"][i]),2)
                            d["Délai d'appro"] =round(float(df["Délai d'appro"][i]),2)
                            d["SAV"] =round(int(df["Quantité"][i])*float(df["SAV"][i]),2)
                            d["COUT TOTAL"]=round((d["COUT MO"]+d["COUT FOURNITURE RATP"]),2)
                            somme=somme+d["COUT TOTAL"]
                            somme1=somme1+d["COUT MO"]
                            #somme2=somme2+d["COUT FOURNITURE RATP"]
                    ddictio['COUT TOTAL']=somme
                    ddictio['COUT MO']=somme1
                    #ddictio["COUT FOURNITURE"]=somme2
                    my_dataframe.append(ddictio)
                    #st.write(my_dataframe)
    dataframe=pd.DataFrame(my_dataframe)
    return dataframe
        
        
def estimation_totale(param,tit):
    placeholder=st.empty()
    if placeholder.button('Estimer') : 
            if tit=='EIFFAGE':
                                                                   st.session_state.data = load_data("BPU_EIFFAGE.xlsx")
            elif tit =='ERI':
                                                                  st.session_state.data = load_data("BPU_ERI.xlsx")
            elif tit =='INEO':
                                                                 st.session_state.data = load_data("BPU_INEO.xlsx")
            elif tit =='SDEL':
                                                                 st.session_state.data = load_data("BPU_SDEL.xlsx") 
            else:
                                                                 st.session_state.data = load_data("BPU_SOGETREL.xlsx")
            placeholder.empty()
            df=pd.read_excel("quantite_projet.xlsx")
            data=st.session_state.data
            ll=[]
            s=df.shape
            params=param
            if (s[0] >0):
                        st.session_state.estimer=True
                        for i in range(len(df)):
                            if param =='PAR DEFAUT':
                                    params=df["Travaux"][i]
                            else:
                                 pass
                            d=dict()
                            d["N°Sous Système"]=df["Sous Système"][i]
                            d["Désignation"]=df["Désignation"][i]
                            d["Travaux"]=df["Travaux"][i]
                            #d["Travaux"]=param
                            d["Quantité"]=df["Quantité"][i]
                            d["CMP préstation unitaire"]=df["CMP"][i]
                            d["Fournitures préstation unitaire"]=df["Fournitures unitaires"][i]
                            lll=(data[data['Désignation']==d["Désignation"]])
                            #print((lll["Prix unitaire MO NUIT COURTE (hors fourniture)|Taux forfaitaire"].unique())[0])
                            df["Taux forfaitaire unitaire JOUR"][i]=(lll["Prix unitaire MO JOUR (hors fourniture)|Taux forfaitaire"].unique())[0]
                            df["Taux forfaitaire unitaire NUIT COURTE"][i]=(lll["Prix unitaire MO NUIT COURTE (hors fourniture)|Taux forfaitaire"].unique())[0]
                            df["Taux forfaitaire unitaire NUIT LONGUE"][i]=(lll["Prix unitaire MO NUIT LONGUE (hors fourniture)|Taux forfaitaire"].unique())[0]
                            
                            df["Fournitures unitaires"][i]=(lll["Fournitures \nP.U en euros"].unique())[0]
                            if (params)== 'JOUR' :
                                d["COUT MO"]=round((int(df["Quantité"][i]))*(float(df["Taux forfaitaire unitaire JOUR"][i])),2)
                            elif (params)== 'NUIT COURTE':
                                d["COUT MO"]=round((int(df["Quantité"][i]))*(float(df["Taux forfaitaire unitaire NUIT COURTE"][i])),2)
                            else:
                                d["COUT MO"]=round((int(df["Quantité"][i]))*(float(df["Taux forfaitaire unitaire NUIT LONGUE"][i])),2)
                            d["COUT FOURNITURE RATP"] =round(int(df["Quantité"][i])*float(df["CMP"][i]),2)
                            d["COUT FOURNITURE TITULAIRE"] =round(int(df["Quantité"][i])*float(df["Fournitures unitaires"][i]),2)
                            d["Durée des travaux"] =round(int(df["Quantité"][i])*float(df["Temps de main d'oeuvre"][i]),2)
                            d["Délai d'appro"] =round(float(df["Délai d'appro"][i]),2)
                            d["SAV"] =round(int(df["Quantité"][i])*float(df["SAV"][i]),2)
                            d["COUT TOTAL"]=round((d["COUT MO"]+d["COUT FOURNITURE RATP"]+d["COUT FOURNITURE TITULAIRE"]),2)
                            ll.append(d)
                        df.to_excel("quantite_projet.xlsx",index=False)
                        st.session_state.qte=df
                        res=pd.DataFrame(ll)
                        st.session_state.estimation1=res
                        tab1, tab2, tab3,tab4,tab5,tab6,tab7= st.tabs(["Estimation générale", "Estimation par préstation", "Estimation par sous système","Visualisations","Equipements utilisés",'Rapport','Synthèse'])
                        with tab1:
                                st.markdown("""
                <style>
                div[data-testid="metric-container"] {
                   background-color: rgba(25,25,112, 1);
                   border: 0px solid rgba(0,128,128,1);
                   padding: 1% 1% 1% 1%;
                   border-radius: 5px;
                   color: rgb(255, 255, 255);
                   overflow-wrap: break-word;
                }
                /* breakline for metric text         */
                div[data-testid="metric-container"] > label[data-testid="stMetricLabel"] > div {
                   overflow-wrap: break-word;
                   white-space: break-spaces;
                   color:White
                }
                </style>
                """, unsafe_allow_html=True)
                                a=round(res["COUT TOTAL"].sum(),2)
                                b=round(res["COUT FOURNITURE RATP"].sum(),2)
                                d=round(res["COUT FOURNITURE TITULAIRE"].sum(),2)
                                c=round(res["COUT MO"].sum(),2)
                                m=round(res["Durée des travaux"].sum(),2)
                                z=round(res["Délai d'appro"].max(),2)
                                st.metric('COUT TOTAL (€)',a)
                                col1,col2,col3=st.columns(3)
                                with col1:
                                      st.metric('COUT DE FOURNITURES RATP (€)',b)
                                with col2:
                                      st.metric('COUT DE FOURNITURES TITULAIRES (€)',d)
                                with col3:
                                      st.metric("COUT DE MAIN D'OEUVRE (€)",c)
                                col1,col2=st.columns(2)   
                                with col2: 
                                    st.metric("COUT EN € DU TITULAIRE : "+str(st.session_state.titu),d+c)
                                    #st.metric("DELAI D'APPROVISIONNEMENT MAXIMAL (Jours)",z)
                                    #st.metric("SAV (€)",rr)
                                
                                
                                
                        with tab2:
                                st.write(res)
                                res.to_excel("Estimation.xlsx",index=False)
                                df_xlsx = to_excell(res)
                                st.download_button(label='📥 Télécharger',
                                                data=df_xlsx ,
                                                file_name= 'ESTIMATION.xlsx')
                        with tab3:
                                my_list=res["N°Sous Système"].unique()
                                dd = st.session_state.syst
                
                                ll=[]
                                for elem in my_list:
                                    d=dict()
                                    my_data=res[res["N°Sous Système"]==elem]
                                    d["N°Sous Système"]=elem
                                    des=dd[dd["N°Sous Système"]==elem]
                                    des=(des["Désignation"].unique())[0]
                                    d["Désignation"]=des
                                    d["COUT FOURNITURE RATP"]=round(my_data["COUT FOURNITURE RATP"].sum(),2)
                                    d["COUT FOURNITURE TITULAIRE"]=round(my_data["COUT FOURNITURE TITULAIRE"].sum(),2)
                                    d["COUT MO"]=round(my_data["COUT MO"].sum(),2)
                                    d["DUREE DE TRAVAUX"]=round(my_data["Durée des travaux"].sum(),2)
                                    d["DELAI D'APPROVISIONNEMENT"]=round(my_data["Délai d'appro"].max(),2)
                                    d["SAV"]=round(my_data["SAV"].sum(),2)
                                    d["COUT TOTAL"]=round(my_data["COUT TOTAL"].sum(),2)
                                    ll.append(d)
                                dataframe=pd.DataFrame(ll)
                                st.write(dataframe)
                                st.session_state.estimation2=dataframe
                                df_xlsx = to_excell(dataframe)
                                st.download_button(label='📥 Télécharger',
                                                data=df_xlsx ,
                                                file_name= 'ESTIMATION-PRESTATION.xlsx')
                                
                                
                        with tab4:
                                    m=get_dataframe(res)
                                    with st.expander("COUT DE FOURNITURE PAR SOUS SYSTEME"):
                                       fig = px.bar(m, x = 'Désignation',y = 'COUT FOURNITURE RATP',title = 'Cout Fourniture par sous système' )
                                       st.plotly_chart(fig)
                                    with st.expander("COUT DE MAIN D'OEUVRE PAR SOUS SYSTEME"):
                                       fig = px.bar(m, x = 'Désignation',y = 'COUT MO',title = 'Cout MO par sous système' )
                                       st.plotly_chart(fig)  
                                       
                                        
                                 
                                    with st.expander("COUT DE FOURNITURE PAR PRESTATION"):
                                       fig = px.bar(res, x = 'Désignation',y = 'COUT FOURNITURE RATP',title = 'Cout Fourniture par préstation' )
                                       st.plotly_chart(fig)
                                    with st.expander("COUT DE MAIN D'OEUVRE PAR PRESTATION"):
                                       fig = px.bar(res, x = 'Désignation',y = 'COUT MO',title = 'Cout MO par préstation' )
                                       st.plotly_chart(fig)
                        with tab5:
                                dat=st.session_state.eq_asoc
                                if dat.shape[0]>0:
                                    dat=dat.fillna(0)
                                    st.dataframe(dat)
                                    df_xlsx = to_excell(dat)
                                    st.download_button(label='📥 Télécharger',data=df_xlsx ,file_name= 'Equipements_quantités.xlsx')
                                else:
                                    st.info('Aucun équipement associé', icon="ℹ️")
                        with tab7:
                             st.markdown(""" <style> .font {
                                font-size:35px ; font-family: 'Cooper Black'; color: #95c2a9;} 
                                </style> """, unsafe_allow_html=True)
                             st.markdown('<p class="font">NUIT COURTE</p>', unsafe_allow_html=True)
                             df1=rapport_synthèse('NUIT COURTE')
                             st.write(df1)
                             st.markdown(""" <style> .font {
                                font-size:35px ; font-family: 'Cooper Black'; color: #95c2a9;} 
                                </style> """, unsafe_allow_html=True)
                             st.markdown('<p class="font">NUIT LONGUE</p>', unsafe_allow_html=True)
                             df2=rapport_synthèse('NUIT LONGUE')
                             st.write(df2)
                             df3=rapport_synthèse('JOUR')
                             st.markdown(""" <style> .font {
                                font-size:35px ; font-family: 'Cooper Black'; color: #95c2a9;} 
                                </style> """, unsafe_allow_html=True)
                             st.markdown('<p class="font">JOUR</p>', unsafe_allow_html=True)
                             st.write(df3)
                             df1['Jour|Nuit']='NUIT COURTE'
                             df2['Jour|Nuit']='NUIT LONGUE'
                             df3['Jour|Nuit']='JOUR'
                             st.markdown(""" <style> .font {
                                font-size:35px ; font-family: 'Cooper Black'; color: #95c2a9;} 
                                </style> """, unsafe_allow_html=True)
                             st.markdown('<p class="font">RECAPITULATIF</p>', unsafe_allow_html=True)
                             df_res=pd.concat([df1, df2])
                             df_res=pd.concat([df_res, df3])
                             df_res['NUIT COURTE-JOUR']=df1['COUT MO']-df3['COUT MO']
                             df_res['NUIT LONGUE-JOUR']=df2['COUT MO']-df3['COUT MO']
                             df_res['NUIT COURTE-NUIT LONGUE']=df1['COUT MO']-df2['COUT MO']
                             st.write(df_res)
                             df_xlsx = to_excell(df_res)
                             st.download_button(label='Télécharger',
                                                data=df_xlsx ,
                                                file_name= 'Synthèse.xlsx')
                             
                        with tab6:   
                               if (st.session_state.estimer):
                                    dictio=st.session_state.projet 
                                    my_tables=[]
                                    my_names=[]
                                    my_dict=dict()
                                    df1=st.session_state.estimation1
                                    a=round(df1["COUT TOTAL"].sum(),2)
                                    b=round(df1["COUT FOURNITURE RATP"].sum(),2)
                                    c=round(df1["COUT MO"].sum(),2)
                                    m=round(df1["Durée des travaux"].sum(),2)
                                    z=round(df1["Délai d'appro"].max(),2)                                            
                                    d=round(df1["COUT FOURNITURE TITULAIRE"].sum(),2)
                                    r=round(df1["SAV"].sum(),2)
                                    dictio['Cout total du projet']=a
                                    dictio['Cout de fourniture  *ratp* du projet']=b
                                    dictio['Cout de fourniture  *titulaire* du projet']=d
                                    dictio["Cout de main d'oeuvre du projet "]=c
                                    dictio["Durée des travaux"]=m
                                    dictio["Délai d'approvisonnement"]=z 
                                    dictio["Total SAV"]=r
                                    st.markdown(""" <style> .font {
                                font-size:35px ; font-family: 'Cooper Black'; color: #95c2a9;} 
                                </style> """, unsafe_allow_html=True)
                                    st.markdown('<p class="font">Estimation générale</p>', unsafe_allow_html=True)
                                    st.write(dictio)
                                    df=st.session_state.qte
                                    my_tables.append(df)
                                    st.markdown(""" <style> .font {
                                font-size:35px ; font-family: 'Cooper Black'; color: #95c2a9;} 
                                </style> """, unsafe_allow_html=True)
                                    st.markdown('<p class="font">Quantités</p>', unsafe_allow_html=True)
                                    st.write(df)
                                    my_names.append('Quantités')
                                    my_tables.append(df1)
                                    st.markdown(""" <style> .font {
                                font-size:35px ; font-family: 'Cooper Black'; color: #95c2a9;} 
                                </style> """, unsafe_allow_html=True)
                                    st.markdown('<p class="font">Estimation par prestation</p>', unsafe_allow_html=True)
                                    st.write(df1)
                                    my_names.append('Estimation par prestation')
                                    df2=st.session_state.estimation2
                                    my_tables.append(df2)
                                    
                                    st.markdown(""" <style> .font {
                                font-size:35px ; font-family: 'Cooper Black'; color: #95c2a9;} 
                                </style> """, unsafe_allow_html=True)
                                    st.markdown('<p class="font">Estimation par sous système</p>', unsafe_allow_html=True)
                                    st.write(df2)
                                    my_names.append('Estimation par sous système')
                                    df5=load_data('Eq_asoc.xlsx')
                                    my_tables.append(df5)
                                    my_names.append('Equipements utilisés')
                                    my_dict=dictio
                                    dictio=pd.DataFrame(dictio.items(), columns=['', ''])
                                    placeholder=st.empty()
                                    doctable(my_tables,my_names,dictio,my_dict,'')
                                          
            else:
                st.error('Aucune quantité saisie!!!') 
                            
    
    
def sauvegarde(dataframe,path):
        dataframe.to_excel('./historique/'+path,index=False)
        st.success('Le chiffrage en cours a été sauvegardé')
        (st.session_state.liste_projet).append(st.session_state.projet)
        ress=pd.DataFrame(st.session_state.liste_projet)
        ress.to_excel('liste_projects.xlsx',index=False)
        
def consulter_sauvegarde():
                        import os
                        d=[]
                        folder_path = "./historique"
                        for path, dirs, files in os.walk(folder_path):
                              for filename in files:
                                   if filename.split('.')[1]=='xlsx':
                                        d.append(filename)
                        if len(d) ==0 :
                            st.error('Aucun chiffrage sauvegardé')
                        else:
                            file=st.selectbox('Sélectionner:',d)
                            dataframe=load_data(folder_path+'/'+file)
                            gb = GridOptionsBuilder.from_dataframe(dataframe)
                            gb.configure_pagination(paginationAutoPageSize=True) #Add pagination
                            gb.configure_side_bar() #Add a sidebar
                            gb.configure_selection('disabled', use_checkbox=True, groupSelectsChildren="Group checkbox select children") 
                            gridOptions = gb.build()
                            grid_response = AgGrid(
                                dataframe,
                                gridOptions=gridOptions,
                                data_return_mode='AS_INPUT', 
                                update_mode='MODEL_CHANGED', 
                                fit_columns_on_grid_load=False,
                                enable_enterprise_modules=True,
                                theme='alpine',
                                height = "800px", 
                                width='100%',
                                reload_data=False
                            )

                            data = grid_response['data']
                            selected = grid_response['selected_rows'] 
                            df = pd.DataFrame(selected)
                            df_xlsx = to_excell(dataframe)
                            st.download_button(label='📥 Télécharger',
                                                data=df_xlsx ,
                                                file_name=file)
def consulter_sauvegarde_p():
                        import os
                        d=[]
                        folder_path = "./historique"
                        for path, dirs, files in os.walk(folder_path):
                              for filename in files:
                                 if filename.split('.')[1]=='xlsx':
                                        d.append(filename)
                        if len(d) ==0 :
                            st.error('Aucun chiffrage sauvegardé')
                        else:
                            file=st.selectbox('Sélectionner:',d)
                            dataframe=load_data(folder_path+'/'+file)
                            df_xlsx = to_excell(dataframe)
                            st.download_button(label='Télécharger',data=df_xlsx ,file_name=file)
import os
def consulter_sauvegarde_prime():
    d=[]
    dataframe=pd.DataFrame()
    folder_path = "./historique"
    for path, dirs, files in os.walk(folder_path):
          for filename in files:
                d.append(filename)
    if len(d)==0:
        st.error("Historique vide")
    else:
          file=st.selectbox('Sélectionner:',d)
          if st.button('Valider'):
                  st.session_state.c_cours=True
                  dataframe=load_data(folder_path+'/'+file)
    return dataframe
   
def check_remove(my_list,num):
    for dico in my_list:
        if dico['N° Projet'] ==num:
            my_list.remove(dico)
    return my_list

def supprimer_fichiers_repertoire(repertoire,fichs):
    # Vérifier si le répertoire existe
    if os.path.exists(repertoire):
        # Obtenir la liste des fichiers dans le répertoire
        fichiers = os.listdir(repertoire)
        if len(fichiers)==0:
            st.error("Historique vide")
        else:
                # Parcourir les fichiers
                for fichier in fichs:
                    if fichier in fichiers:
                             
                                      # Construire le chemin absolu du fichier
                                      chemin_fichier = os.path.join(repertoire, fichier)
                    # Vérifier si le chemin correspond à un fichier (plutôt qu'à un sous-répertoire)
                    if os.path.isfile(chemin_fichier):
                        # Supprimer le fichier
                        num_projet=(fichier.split('_')[0])
                        os.remove(chemin_fichier)
                        st.success(f"Le fichier {fichier} a été supprimé.")
                        (st.session_state.liste_projet)=check_remove(st.session_state.liste_projet,num_projet)
                        ress=pd.DataFrame(st.session_state.liste_projet)
                        ress.to_excel('liste_projects.xlsx',index=False)
    else:
        st.error(f"Le répertoire {repertoire} n'existe pas.")

def vider_hist(repertoire):
    # Vérifier si le répertoire existe
    if os.path.exists(repertoire):
        # Obtenir la liste des fichiers dans le répertoire
        fichiers = os.listdir(repertoire)
        #print(fichiers)
        if len(fichiers)<2:
            st.error('Historique déja vide')
            
        else:
            # Parcourir les fichiers
            for fichier in fichiers:
                        # Construire le chemin absolu du fichier
                chemin_fichier = os.path.join(repertoire, fichier)
                # Vérifier si le chemin correspond à un fichier (plutôt qu'à un sous-répertoire)
                if os.path.isfile(chemin_fichier):
                    # Supprimer le fichier
                    os.remove(chemin_fichier)
            st.success("L'historique a été vidé")
            (st.session_state.liste_projet)=[]
            ress=pd.DataFrame(st.session_state.liste_projet)
            ress.to_excel('liste_projects.xlsx',index=False)
    else:
        st.error(f"Le répertoire {repertoire} n'existe pas.")
            
def df_to_word(df, doc) -> docx.Document:
    table = doc.add_table(df.shape[0]+1, df.shape[1])
    for j in range(df.shape[-1]):
        table.cell(0,j).text = df.columns[j]
    for i in range(df.shape[0]):
        for j in range(df.shape[-1]):
            table.cell(i+1,j).text = str(df.values[i,j])

    return doc  
    


    
def doctable(datas, tabletitles, d,my_dict,pathfile):
    from docx import Document
    from docx.shared import Pt, Mm
    document = Document()
    section = document.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(297)
    section.left_margin = Mm(20)
    section.right_margin = Mm(20)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.header_distance = Mm(12.7)
    section.footer_distance = Mm(12.7)
    document.add_picture('R.png',width=Inches(2))
    para = document.add_paragraph().add_run('Informations du projet :')
    para.font.color.rgb = RGBColor(0x22, 0x8b, 0x22)
    para.font.size = Pt(16)
    document=df_to_word(d,document)
    para1 = document.add_paragraph().add_run('Estimations à partir des quantités projet')
    para1.font.color.rgb = RGBColor(0x22, 0x8b, 0x22)
    para1.font.size = Pt(16)
    for data ,tabletitle in zip(datas,tabletitles):
            document.add_heading(tabletitle)
            document=df_to_word(data,document)
    document.add_page_break()
    #document.save(pathfile)
    my_path='rapport_'+str(my_dict['N° Projet'])+'_'+str(my_dict['Date'])+'_'+str(st.session_state.titu)+'.docx'
    #document.save(my_path)
    import io
    bio = io.BytesIO()
    document.save(bio)
    st.download_button(label="Télécharger",data=bio.getvalue(),file_name=my_path,mime="docx"
        )
    
    return 0
